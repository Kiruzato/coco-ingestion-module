# ==============================================================================
# WARNING: SYNCHRONIZED FILE
# ==============================================================================
# This file exists in TWO locations:
#   1. WEB_APP/modules/document_manager.py       (runtime)
#   2. INGESTION_MODULE/modules/document_manager.py  (ingestion)
#
# These copies are intentionally independent to maintain strict architectural
# isolation between WEB_APP and INGESTION_MODULE. Neither module may import
# from the other.
#
# ANY CHANGES to this file MUST be manually mirrored to the other copy.
# Failure to synchronize will cause behavioral divergence between runtime
# and ingestion environments.
# ==============================================================================
"""
Document Management System - Phase 3 + Text Normalization
==========================================================
This module handles document ingestion, metadata management, and indexing for the
campus RAG chatbot. Supports PDF, DOCX, and TXT formats.

Key Features:
- Multi-format document loading (PDF, DOCX, TXT)
- Enhanced metadata tracking (document_id, timestamp, file type)
- Duplicate detection
- Incremental ingestion
- Document registry management
- Text normalization for consistent retrieval (case-insensitive matching)
"""

import os
import json
import hashlib
from pathlib import Path
from typing import List, Dict, Optional, Tuple
from datetime import datetime
import logging
from dotenv import load_dotenv

# Load environment variables
load_dotenv()

# Configure logging early (before imports that may log warnings)
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

# Document loaders
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain_openai import OpenAIEmbeddings
from langchain_core.documents import Document

# PDF loader
from pypdf import PdfReader

# Layout-aware PDF parsing (Phase 18)
from unstructured.partition.pdf import partition_pdf

# DOCX loader
from docx import Document as DocxDocument

# Text normalization for consistent retrieval
from modules.text_normalizer import normalize_text


# ==============================================================================
# CONFIGURATION
# ==============================================================================

OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
SUPPORTED_FORMATS = ['.txt', '.pdf', '.docx']


# ==============================================================================
# METADATA SCHEMA (Phase 23)
# ==============================================================================

# Required fields for ALL chunks
REQUIRED_METADATA = {
    'document_id': str,
    'document_name': str,
    'file_type': str,
    'ingestion_timestamp': str,
    'chunk_id': int,
    'section': str,
    'original_text': str,
}

# Additional required fields for PDF chunks with layout-aware parsing
REQUIRED_PDF_LAYOUT_METADATA = {
    'section_title': str,
    'element_types': list,
    'page_numbers': list,
}

# Additional required fields for synthetic chunks
REQUIRED_SYNTHETIC_METADATA = {
    'is_synthetic': bool,
    'entity_type': str,
    'source_chunk_ids': list,
}

# Suspicious values that trigger warnings
SUSPICIOUS_VALUES = {
    'section': ['General Information'],
    'section_title': ['General Information', 'Document Content'],
}


# ==============================================================================
# DOCUMENT REGISTRY MANAGEMENT
# ==============================================================================

class DocumentRegistry:
    """
    Manages a registry of ingested documents to track metadata and prevent duplicates.
    Registry is stored as a JSON file.
    """

    def __init__(self, registry_path: Path):
        """
        Initialize the document registry.

        Args:
            registry_path: Path to the JSON registry file
        """
        self.registry_path = registry_path
        self.documents = self._load_registry()

    def _load_registry(self) -> Dict:
        """Load the registry from disk or create a new one."""
        if self.registry_path.exists():
            with open(self.registry_path, 'r', encoding='utf-8') as f:
                return json.load(f)
        return {}

    def _save_registry(self):
        """Save the registry to disk."""
        with open(self.registry_path, 'w', encoding='utf-8') as f:
            json.dump(self.documents, f, indent=2, default=str)
        logger.info(f"Registry saved to {self.registry_path}")

    def add_document(self, document_id: str, metadata: Dict):
        """
        Add a document to the registry.

        Args:
            document_id: Unique identifier for the document
            metadata: Document metadata dictionary
        """
        self.documents[document_id] = metadata
        self._save_registry()
        logger.info(f"Added document to registry: {metadata['document_name']} (ID: {document_id})")

    def remove_document(self, document_id: str) -> bool:
        """
        Remove a document from the registry.

        Args:
            document_id: Document identifier to remove

        Returns:
            True if removed, False if not found
        """
        if document_id in self.documents:
            doc_name = self.documents[document_id].get('document_name', 'Unknown')
            del self.documents[document_id]
            self._save_registry()
            logger.info(f"Removed document from registry: {doc_name} (ID: {document_id})")
            return True
        return False

    def get_document(self, document_id: str) -> Optional[Dict]:
        """Get document metadata by ID."""
        return self.documents.get(document_id)

    def list_documents(self) -> List[Dict]:
        """
        List all documents in the registry.

        Returns:
            List of document metadata dictionaries
        """
        return [
            {"document_id": doc_id, **metadata}
            for doc_id, metadata in self.documents.items()
        ]

    def document_exists(self, file_hash: str) -> bool:
        """
        Check if a document with the given file hash already exists.

        Args:
            file_hash: SHA256 hash of the file

        Returns:
            True if document exists, False otherwise
        """
        for doc_metadata in self.documents.values():
            if doc_metadata.get('file_hash') == file_hash:
                return True
        return False

    def get_by_hash(self, file_hash: str) -> Optional[Tuple[str, Dict]]:
        """
        Get document ID and metadata by file hash.

        Returns:
            Tuple of (document_id, metadata) or None
        """
        for doc_id, metadata in self.documents.items():
            if metadata.get('file_hash') == file_hash:
                return (doc_id, metadata)
        return None


# ==============================================================================
# DOCUMENT LOADING UTILITIES
# ==============================================================================

def calculate_file_hash(file_path: Path) -> str:
    """
    Calculate SHA256 hash of a file for duplicate detection.

    Args:
        file_path: Path to the file

    Returns:
        SHA256 hash as hexadecimal string
    """
    sha256_hash = hashlib.sha256()
    with open(file_path, "rb") as f:
        for byte_block in iter(lambda: f.read(4096), b""):
            sha256_hash.update(byte_block)
    return sha256_hash.hexdigest()


def extract_section_name(text_chunk: str) -> str:
    """
    Extract section name from a text chunk.

    Args:
        text_chunk: The text chunk to analyze

    Returns:
        Section name or "General Information"
    """
    lines = text_chunk.split('\n')
    for line in lines[:3]:
        line = line.strip()
        if line and (line.isupper() or (line[0].isupper() and len(line.split()) <= 5)):
            return line
    return "General Information"


def is_appendix_section(section_title: str) -> tuple:
    """
    Phase 22: Detect if a section is an appendix and extract its identifier.

    Appendix sections get special handling:
    - Larger max_size to keep content intact
    - Never merged with non-appendix content
    - Marked with is_appendix metadata

    Args:
        section_title: The section title to check

    Returns:
        Tuple of (is_appendix: bool, appendix_id: str or None)

    Examples:
        "Appendix A" -> (True, "A")
        "Appendix D - Prayer" -> (True, "D")
        "APPENDIX C" -> (True, "C")
        "Columban Hymn" -> (True, "CONTENT")
        "Academic Programs" -> (False, None)
    """
    import re
    title_lower = section_title.lower().strip()

    # Pattern 1: "appendix" followed by letter/number
    match = re.match(r'^appendix\s*([a-z0-9])?', title_lower)
    if match:
        appendix_id = match.group(1).upper() if match.group(1) else "UNKNOWN"
        logger.info(f"[PHASE22] Detected appendix section: '{section_title}' (ID: {appendix_id})")
        return True, appendix_id

    # Pattern 2: Section title that IS an appendix item (Hymn, Prayer)
    appendix_items = ['columban hymn', 'prayer to st. columban', 'prayer to st columban']
    if any(item in title_lower for item in appendix_items):
        logger.info(f"[PHASE22] Detected appendix content section: '{section_title}'")
        return True, "CONTENT"

    return False, None


# ==============================================================================
# METADATA VALIDATION (Phase 23)
# ==============================================================================

def validate_chunk_metadata(chunk: Document, is_pdf_layout: bool = False) -> tuple:
    """
    Phase 23: Validate chunk metadata against schema.

    Args:
        chunk: Document object with metadata
        is_pdf_layout: Whether this chunk was created via layout-aware parsing

    Returns:
        Tuple of (is_valid: bool, errors: List[str], warnings: List[str])
    """
    errors = []
    warnings = []
    metadata = chunk.metadata
    chunk_id = metadata.get('chunk_id', 'UNKNOWN')

    # Check required fields
    for field, expected_type in REQUIRED_METADATA.items():
        if field not in metadata:
            errors.append(f"Chunk {chunk_id}: Missing required field '{field}'")
        elif not isinstance(metadata[field], expected_type):
            errors.append(f"Chunk {chunk_id}: Field '{field}' wrong type "
                         f"(expected {expected_type.__name__}, got {type(metadata[field]).__name__})")

    # Check PDF layout-specific fields
    if is_pdf_layout:
        for field, expected_type in REQUIRED_PDF_LAYOUT_METADATA.items():
            if field not in metadata:
                warnings.append(f"Chunk {chunk_id}: Missing PDF layout field '{field}'")
            elif not isinstance(metadata[field], expected_type):
                errors.append(f"Chunk {chunk_id}: Field '{field}' wrong type")

    # Check synthetic chunk fields
    if metadata.get('is_synthetic', False):
        for field, expected_type in REQUIRED_SYNTHETIC_METADATA.items():
            if field not in metadata:
                errors.append(f"Chunk {chunk_id}: Synthetic chunk missing '{field}'")
            elif not isinstance(metadata[field], expected_type):
                errors.append(f"Chunk {chunk_id}: Field '{field}' wrong type")

    # Check for suspicious values
    for field, suspicious_list in SUSPICIOUS_VALUES.items():
        if field in metadata and metadata[field] in suspicious_list:
            warnings.append(f"Chunk {chunk_id}: Suspicious value '{metadata[field]}'")

    # Validate page_numbers elements are integers
    if 'page_numbers' in metadata and isinstance(metadata['page_numbers'], list):
        for pn in metadata['page_numbers']:
            if not isinstance(pn, int):
                errors.append(f"Chunk {chunk_id}: page_numbers contains non-integer: {pn}")

    # Validate element_types elements are strings
    if 'element_types' in metadata and isinstance(metadata['element_types'], list):
        for et in metadata['element_types']:
            if not isinstance(et, str):
                errors.append(f"Chunk {chunk_id}: element_types contains non-string: {et}")

    is_valid = len(errors) == 0
    return is_valid, errors, warnings


def validate_all_chunks(documents: List[Document], is_pdf_layout: bool = False) -> tuple:
    """
    Phase 23: Validate all chunks and log results.

    Args:
        documents: List of Document objects to validate
        is_pdf_layout: Whether chunks were created via layout-aware parsing

    Returns:
        Tuple of (all_valid: bool, total_errors: int, total_warnings: int)
    """
    all_errors = []
    all_warnings = []

    for doc in documents:
        is_valid, errors, warnings = validate_chunk_metadata(doc, is_pdf_layout)
        all_errors.extend(errors)
        all_warnings.extend(warnings)

    # Log validation results
    if all_errors:
        logger.error(f"[PHASE23] Metadata validation: {len(all_errors)} errors")
        for error in all_errors[:10]:
            logger.error(f"  - {error}")
        if len(all_errors) > 10:
            logger.error(f"  ... and {len(all_errors) - 10} more errors")

    if all_warnings:
        logger.warning(f"[PHASE23] Metadata validation: {len(all_warnings)} warnings")
        for warning in all_warnings[:10]:
            logger.warning(f"  - {warning}")
        if len(all_warnings) > 10:
            logger.warning(f"  ... and {len(all_warnings) - 10} more warnings")

    return len(all_errors) == 0, len(all_errors), len(all_warnings)


def log_metadata_statistics(documents: List[Document]):
    """
    Phase 23: Log metadata statistics for observability.

    Args:
        documents: List of Document objects
    """
    if not documents:
        logger.info("[PHASE23] No documents to analyze")
        return

    stats = {
        'total': len(documents),
        'synthetic': sum(1 for d in documents if d.metadata.get('is_synthetic')),
        'appendix': sum(1 for d in documents if d.metadata.get('is_appendix')),
        'sections': set(),
        'element_types': set(),
        'general_info': 0,
        'sizes': [],
    }

    for doc in documents:
        meta = doc.metadata

        # Collect sections
        section = meta.get('section_title') or meta.get('section', '')
        if section:
            stats['sections'].add(section)
        if section == 'General Information':
            stats['general_info'] += 1

        # Collect element types
        for et in meta.get('element_types', []):
            stats['element_types'].add(et)

        # Chunk sizes
        stats['sizes'].append(len(doc.page_content))

    # Calculate averages
    avg_size = sum(stats['sizes']) / len(stats['sizes']) if stats['sizes'] else 0
    min_size = min(stats['sizes']) if stats['sizes'] else 0
    max_size = max(stats['sizes']) if stats['sizes'] else 0

    # Log statistics
    logger.info("[PHASE23] ========== METADATA STATISTICS ==========")
    logger.info(f"[PHASE23] Total chunks: {stats['total']}")
    logger.info(f"[PHASE23] Synthetic chunks: {stats['synthetic']}")
    logger.info(f"[PHASE23] Appendix chunks: {stats['appendix']}")
    logger.info(f"[PHASE23] Unique sections: {len(stats['sections'])}")
    logger.info(f"[PHASE23] Element types: {sorted(stats['element_types'])}")
    logger.info(f"[PHASE23] Chunk size: avg={avg_size:.0f}, min={min_size}, max={max_size}")
    if stats['general_info'] > 0:
        logger.warning(f"[PHASE23] 'General Information' sections: {stats['general_info']} (detection gaps)")
    logger.info("[PHASE23] ============================================")


def load_txt_document(file_path: Path) -> str:
    """
    Load a text document.

    Args:
        file_path: Path to the TXT file

    Returns:
        Document text content
    """
    with open(file_path, 'r', encoding='utf-8') as f:
        return f.read()


def load_pdf_document(file_path: Path) -> str:
    """
    Load a PDF document.

    Args:
        file_path: Path to the PDF file

    Returns:
        Extracted text content

    Raises:
        ImportError: If PDF loader is not available
        Exception: If PDF cannot be read
    """
    try:
        reader = PdfReader(str(file_path))
        text_content = []

        for page_num, page in enumerate(reader.pages, 1):
            text = page.extract_text()
            if text.strip():
                text_content.append(text)

        logger.info(f"Extracted text from {len(reader.pages)} pages in {file_path.name}")
        return "\n\n".join(text_content)

    except Exception as e:
        logger.error(f"Failed to load PDF {file_path.name}: {str(e)}")
        raise


def load_pdf_document_layout_aware(file_path: Path) -> List[Dict]:
    """
    Load a PDF document with layout-aware parsing (Phase 18).

    Extracts structured elements (Title, NarrativeText, ListItem, Table) with metadata.

    Args:
        file_path: Path to the PDF file

    Returns:
        List of element dictionaries with:
        - type: Element type (Title, NarrativeText, ListItem, Table, etc.)
        - text: Element text content
        - metadata: Dict with page_number, etc.

    Raises:
        ImportError: If layout-aware parser is not available
        Exception: If PDF cannot be parsed (caller should fallback to linear extraction)
    """
    try:
        # Partition PDF using fast strategy (suitable for most campus documents)
        # strategy="fast" uses rule-based parsing without OCR for better performance
        elements = partition_pdf(
            filename=str(file_path),
            strategy="fast",
            infer_table_structure=True  # Extract tables as structured data
        )

        logger.info(f"Layout-aware parsing found {len(elements)} elements in {file_path.name}")

        # Convert elements to dictionaries with standardized structure
        extracted_elements = []
        element_type_counts = {}

        for elem in elements:
            elem_type = elem.category if hasattr(elem, 'category') else type(elem).__name__
            
            # Count element types for logging
            element_type_counts[elem_type] = element_type_counts.get(elem_type, 0) + 1

            # Extract text content
            # For tables, prefer HTML representation if available
            if elem_type == "Table" and hasattr(elem, 'metadata'):
                text_content = getattr(elem.metadata, 'text_as_html', None) or str(elem)
            else:
                text_content = str(elem)

            # Extract metadata
            elem_metadata = {}
            if hasattr(elem, 'metadata'):
                # Capture page number if available
                if hasattr(elem.metadata, 'page_number'):
                    elem_metadata['page_number'] = elem.metadata.page_number

            extracted_elements.append({
                'type': elem_type,
                'text': text_content,
                'metadata': elem_metadata
            })

        # Log element type distribution
        type_summary = ', '.join([f"{typ}: {count}" for typ, count in element_type_counts.items()])
        logger.info(f"Element types in {file_path.name}: {type_summary}")

        return extracted_elements

    except Exception as e:
        logger.warning(f"Layout-aware parsing failed for {file_path.name}: {str(e)}")
        raise  # Caller will handle fallback to linear extraction


def group_elements_by_section(elements: List[Dict]) -> List[Dict]:
    """
    Group extracted PDF elements into logical sections based on Title elements (Phase 18).

    Args:
        elements: List of element dictionaries from load_pdf_document_layout_aware

    Returns:
        List of section dictionaries with:
        - section_title: Section heading (or "Document Content" if no title)
        - elements: List of element dicts belonging to this section
        - element_types: Set of element types present in section
        - page_numbers: Set of page numbers spanned by section
        - is_appendix: Boolean flag for appendix sections (Phase 22)
        - appendix_id: Appendix identifier if applicable (Phase 22)
    """
    sections = []
    current_section = {
        'section_title': 'Document Content',
        'elements': [],
        'element_types': set(),
        'page_numbers': set(),
        'is_appendix': False,
        'appendix_id': None
    }

    for elem in elements:
        elem_type = elem['type']

        # Title elements start a new section
        if elem_type == 'Title':
            # Save previous section if it has content
            if current_section['elements']:
                sections.append(current_section)

            # Phase 22: Check if this is an appendix section
            section_title = elem['text'].strip()
            is_appendix, appendix_id = is_appendix_section(section_title)

            # Start new section with this title
            current_section = {
                'section_title': section_title,
                'elements': [elem],
                'element_types': {elem_type},
                'page_numbers': {elem['metadata'].get('page_number')} if elem['metadata'].get('page_number') else set(),
                'is_appendix': is_appendix,
                'appendix_id': appendix_id
            }
        else:
            # Add element to current section
            current_section['elements'].append(elem)
            current_section['element_types'].add(elem_type)
            if elem['metadata'].get('page_number'):
                current_section['page_numbers'].add(elem['metadata'].get('page_number'))

    # Don't forget the last section
    if current_section['elements']:
        sections.append(current_section)

    logger.info(f"Grouped {len(elements)} elements into {len(sections)} sections")
    
    # Phase 18 Fix: Merge related administrative sections
    sections = merge_related_admin_sections(sections)
    
    return sections


def merge_related_admin_sections(sections: List[Dict]) -> List[Dict]:
    """
    Merge related administrative sections with HARD BOUNDARIES for enumeration groups.

    Phase 18 Structural Fix: Treats "Deans" as a protected enumeration group that must
    never be mixed with other administrative roles (Directors, VPs, Chairs). This ensures
    deterministic enumeration completeness for queries like "Who are the deans".

    Phase 22: Appendix sections are never merged with non-appendix content.

    Design principles:
    1. Enumeration groups (Deans, Directors, etc.) are HARD BOUNDARIES
    2. Never merge across different role types
    3. Collect ALL instances of same role type into ONE canonical group
    4. Enumerated entities appear FIRST in chunk content
    5. Deterministic, not ranking-dependent
    6. Appendix sections stay isolated (Phase 22)

    Args:
        sections: List of section dictionaries from group_elements_by_section

    Returns:
        List of sections with enumeration groups properly isolated
    """
    if not sections:
        return sections

    # Phase 22: Separate appendix sections - they are never merged
    appendix_sections = []
    non_appendix_sections = []

    for section in sections:
        if section.get('is_appendix', False):
            appendix_sections.append(section)
            logger.info(f"[PHASE22] Appendix section isolated: '{section['section_title']}'")
        else:
            non_appendix_sections.append(section)

    # If no non-appendix sections, just return appendix sections
    if not non_appendix_sections:
        return appendix_sections

    # Apply existing merge logic only to non-appendix sections
    sections = non_appendix_sections
    
    # Define enumeration groups (hard boundaries - never mix)
    ENUMERATION_GROUPS = {
        'deans': ['dean'],  # Exact match for dean roles
        'directors': ['director'],  # Director roles
        'vps': ['vice president', 'vp'],  # VP roles
        'chairs': ['chair', 'chairperson'],  # Department chairs
    }
    
    def classify_section(section: Dict) -> tuple:
        """
        Classify section into enumeration group.
        Returns: (group_name, is_enumeration_section, contains_person_names, is_title_only)
        """
        title_lower = section['section_title'].lower().strip()
        section_text = ' '.join([elem['text'] for elem in section['elements']]).lower()
        
        # Check if this is a title-only section (just the heading, no content)
        is_title_only = (len(section['elements']) == 1 and 
                        section['elements'][0]['type'] == 'Title')
        
        # Check for person name indicators
        has_person_names = any(
            title in section_text 
            for title in ['dr.', 'engr.', 'arch.', 'prof.', 'mr.', 'ms.', 'mrs.']
        )
        
        # Check for exact enumeration group matches
        for group_name, keywords in ENUMERATION_GROUPS.items():
            for keyword in keywords:
                # Title match (high confidence)
                if title_lower == keyword or title_lower == keyword + 's':
                    return (group_name, True, has_person_names, is_title_only)
                
                # Content match with person names (enumeration pattern)
                if keyword in section_text and has_person_names:
                    # Additional check: multiple instances suggest enumeration
                    if section_text.count(keyword) >= 1:
                        return (group_name, True, has_person_names, is_title_only)
        
        return (None, False, has_person_names, is_title_only)
    
    # First pass: Identify and group enumeration sections
    # Handle title-only sections by merging with following section
    enumeration_groups = {group: [] for group in ENUMERATION_GROUPS.keys()}
    other_sections = []
    section_classifications = []
    
    i = 0
    while i < len(sections):
        section = sections[i]
        group_name, is_enum, has_names, is_title_only = classify_section(section)
        section_classifications.append((group_name, is_enum, has_names, is_title_only))
        
        # Handle title-only enumeration sections
        if is_enum and is_title_only and group_name:
            # This is a title-only section like "Deans" with no content
            # Merge it with the next section(s) that contain the actual data
            merged_section = {
                'section_title': section['section_title'],  # Keep the clean title
                'elements': list(section['elements']),
                'element_types': set(section['element_types']),
                'page_numbers': set(section['page_numbers']),
                'is_appendix': False,  # Phase 22: Non-appendix by definition
                'appendix_id': None
            }
            
            # Look ahead for content sections
            j = i + 1
            while j < len(sections):
                next_section = sections[j]
                next_group, next_is_enum, next_has_names, next_is_title_only = classify_section(next_section)
                
                # Merge if:
                # 1. Next section has person names (the actual enumeration data)
                # 2. Same or adjacent pages
                # 3. Not another title-only section
                if next_has_names and not next_is_title_only:
                    next_pages = next_section['page_numbers']
                    current_pages = merged_section['page_numbers']
                    
                    if current_pages and next_pages:
                        max_current = max(current_pages)
                        min_next = min(next_pages)
                        
                        # Merge if on same page or adjacent
                        if abs(min_next - max_current) <= 1:
                            merged_section['elements'].extend(next_section['elements'])
                            merged_section['element_types'].update(next_section['element_types'])
                            merged_section['page_numbers'].update(next_section['page_numbers'])
                            j += 1
                            continue
                
                # Stop merging
                break
            
            enumeration_groups[group_name].append((i, merged_section))
            i = j  # Skip merged sections
        elif is_enum and group_name:
            enumeration_groups[group_name].append((i, section))
            i += 1
        else:
            other_sections.append((i, section))
            i += 1
    
    # Second pass: Merge enumeration groups
    merged_sections = []
    processed_indices = set()
    
    for group_name, group_sections in enumeration_groups.items():
        if not group_sections:
            continue
        
        # Collect all sections for this enumeration group
        # Group by page proximity (within 3 pages = same logical block)
        page_clusters = []
        
        for idx, section in sorted(group_sections, key=lambda x: min(x[1]['page_numbers']) if x[1]['page_numbers'] else 0):
            section_pages = section['page_numbers']
            if not section_pages:
                continue
            
            # Find cluster to add to
            added = False
            for cluster in page_clusters:
                cluster_pages = set()
                for _, s in cluster:
                    cluster_pages.update(s['page_numbers'])
                
                # Check if within 3 pages of cluster
                if cluster_pages and section_pages:
                    min_section = min(section_pages)
                    max_section = max(section_pages)
                    min_cluster = min(cluster_pages)
                    max_cluster = max(cluster_pages)
                    
                    if (abs(min_section - max_cluster) <= 3 or 
                        abs(max_section - min_cluster) <= 3 or
                        section_pages.intersection(cluster_pages)):
                        cluster.append((idx, section))
                        added = True
                        break
            
            if not added:
                page_clusters.append([(idx, section)])
        
        # Create merged sections for each cluster
        for cluster in page_clusters:
            if not cluster:
                continue
            
            # Sort by page number to maintain document order
            cluster.sort(key=lambda x: min(x[1]['page_numbers']) if x[1]['page_numbers'] else 0)
            
            # Create merged section with enumeration-first ordering
            merged_elements = []
            merged_pages = set()
            merged_types = set()
            
            # CRITICAL: Collect elements in enumeration-first order
            # 1. First, collect all Title elements (section headers)
            # 2. Then, collect all person name elements (the actual enumeration)
            # 3. Finally, collect supporting text
            
            title_elements = []
            person_elements = []
            other_elements = []
            
            for idx, section in cluster:
                processed_indices.add(idx)
                merged_pages.update(section['page_numbers'])
                merged_types.update(section['element_types'])
                
                for elem in section['elements']:
                    elem_text = elem['text'].lower()
                    elem_type = elem['type']
                    
                    # Classify element
                    if elem_type == 'Title':
                        # Check if it's the main enumeration title
                        if any(kw in elem_text for kw in ENUMERATION_GROUPS[group_name]):
                            title_elements.insert(0, elem)  # Main title first
                        else:
                            title_elements.append(elem)
                    elif any(title in elem_text for title in ['dr.', 'engr.', 'arch.', 'prof.', 'mr.', 'ms.', 'mrs.']):
                        # Person name element - this is the enumeration data
                        person_elements.append(elem)
                    else:
                        other_elements.append(elem)
            
            # Assemble in correct order: Titles → People → Other
            merged_elements = title_elements + person_elements + other_elements
            
            # Determine canonical title
            canonical_title = group_name.capitalize()  # e.g., "Deans", "Directors"
            
            # Check if any section has the exact canonical title
            for idx, section in cluster:
                if section['section_title'].lower().strip() in [group_name, group_name + 's']:
                    canonical_title = section['section_title']
                    break
            
            merged_section = {
                'section_title': canonical_title,
                'elements': merged_elements,
                'element_types': merged_types,
                'page_numbers': merged_pages,
                'is_appendix': False,  # Phase 22: Non-appendix by definition
                'appendix_id': None
            }

            merged_sections.append((min(merged_pages) if merged_pages else 0, merged_section))
    
    # Add non-enumeration sections
    for idx, section in other_sections:
        if idx not in processed_indices:
            merged_sections.append((min(section['page_numbers']) if section['page_numbers'] else idx, section))
            processed_indices.add(idx)
    
    # Sort by page number to maintain document order
    merged_sections.sort(key=lambda x: x[0])
    result = [section for _, section in merged_sections]

    if len(result) < len(sections):
        logger.info(f"Merged {len(sections)} sections into {len(result)} (isolated {len(sections) - len(result)} enumeration groups)")

    # Phase 22: Append appendix sections at the end (they were isolated earlier)
    if appendix_sections:
        # Sort appendix sections by page number
        appendix_sections.sort(key=lambda s: min(s['page_numbers']) if s['page_numbers'] else float('inf'))
        result.extend(appendix_sections)
        logger.info(f"[PHASE22] Appended {len(appendix_sections)} appendix sections to result")

    return result




def create_chunks_from_sections(
    sections: List[Dict],
    document_id: str,
    document_name: str,
    file_type: str,
    ingestion_timestamp: str,
    target_size: int = 500,
    max_size: int = 800
) -> List[Document]:
    """
    Create chunks from sections with smart size management (Phase 18).

    Strategy:
    - Keep small sections intact (< max_size)
    - Split large sections at element boundaries
    - Preserve lists and tables together when possible
    - Enrich metadata with section context

    Args:
        sections: List of section dictionaries from group_elements_by_section
        document_id: Unique document identifier
        document_name: Original filename
        file_type: File extension
        ingestion_timestamp: ISO format timestamp
        target_size: Target chunk size in characters
        max_size: Maximum chunk size before forced split

    Returns:
        List of Document objects with enriched metadata
    """
    documents = []
    chunk_id = 0

    # Phase 22: Configuration for appendix sections
    APPENDIX_MAX_SIZE = 2000  # Keep appendix content intact if under 2000 chars

    for section in sections:
        section_title = section['section_title']
        section_elements = section['elements']
        section_element_types = list(section['element_types'])  # Convert set to list
        section_page_numbers = sorted(list(section['page_numbers']))  # Convert set to sorted list

        # Phase 22: Extract appendix flags
        is_appendix = section.get('is_appendix', False)
        appendix_id = section.get('appendix_id')

        # Phase 22: Use larger max_size for appendix sections to keep them intact
        if is_appendix:
            effective_max_size = APPENDIX_MAX_SIZE
            logger.info(f"[PHASE22] Appendix section '{section_title}' using max_size={effective_max_size}")
        else:
            effective_max_size = max_size

        # Combine all element texts in this section
        section_text = '\n\n'.join([elem['text'] for elem in section_elements])
        section_length = len(section_text)

        # Case 1: Section fits within effective_max_size - keep intact
        if section_length <= effective_max_size:
            normalized_text = normalize_text(section_text)

            metadata = {
                "document_id": document_id,
                "document_name": document_name,
                "file_type": file_type,
                "ingestion_timestamp": ingestion_timestamp,
                "chunk_id": chunk_id,
                "section": extract_section_name(section_text),  # Legacy compatibility
                "section_title": section_title,  # Phase 18: Explicit section title
                "element_types": section_element_types,  # Phase 18: Element types in chunk
                "page_numbers": section_page_numbers,  # Phase 18: Pages spanned
                "original_text": section_text,
                "is_appendix": is_appendix,  # Phase 22: Appendix flag
                "appendix_id": appendix_id   # Phase 22: Appendix identifier
            }

            doc = Document(page_content=normalized_text, metadata=metadata)
            documents.append(doc)
            chunk_id += 1

        # Case 2: Section too large - split at element boundaries
        else:
            current_chunk_elements = []
            current_chunk_length = 0

            for elem in section_elements:
                elem_text = elem['text']
                elem_length = len(elem_text)

                # If adding this element exceeds effective_max_size and we have content, save current chunk
                if current_chunk_length + elem_length > effective_max_size and current_chunk_elements:
                    # Save current chunk
                    chunk_text = '\n\n'.join([e['text'] for e in current_chunk_elements])
                    normalized_text = normalize_text(chunk_text)

                    # Aggregate page numbers from elements in this chunk
                    chunk_page_numbers = sorted(list(set([
                        e['metadata'].get('page_number')
                        for e in current_chunk_elements
                        if e['metadata'].get('page_number')
                    ])))

                    # Aggregate element types
                    chunk_element_types = list(set([e['type'] for e in current_chunk_elements]))

                    metadata = {
                        "document_id": document_id,
                        "document_name": document_name,
                        "file_type": file_type,
                        "ingestion_timestamp": ingestion_timestamp,
                        "chunk_id": chunk_id,
                        "section": extract_section_name(chunk_text),
                        "section_title": section_title,
                        "element_types": chunk_element_types,
                        "page_numbers": chunk_page_numbers,
                        "original_text": chunk_text,
                        "is_appendix": is_appendix,  # Phase 22
                        "appendix_id": appendix_id   # Phase 22
                    }

                    doc = Document(page_content=normalized_text, metadata=metadata)
                    documents.append(doc)
                    chunk_id += 1

                    # Reset for next chunk
                    current_chunk_elements = []
                    current_chunk_length = 0

                # Add element to current chunk
                current_chunk_elements.append(elem)
                current_chunk_length += elem_length + 2  # +2 for '\n\n' separator

            # Don't forget remaining elements
            if current_chunk_elements:
                chunk_text = '\n\n'.join([e['text'] for e in current_chunk_elements])
                normalized_text = normalize_text(chunk_text)

                chunk_page_numbers = sorted(list(set([
                    e['metadata'].get('page_number')
                    for e in current_chunk_elements
                    if e['metadata'].get('page_number')
                ])))

                chunk_element_types = list(set([e['type'] for e in current_chunk_elements]))

                metadata = {
                    "document_id": document_id,
                    "document_name": document_name,
                    "file_type": file_type,
                    "ingestion_timestamp": ingestion_timestamp,
                    "chunk_id": chunk_id,
                    "section": extract_section_name(chunk_text),
                    "section_title": section_title,
                    "element_types": chunk_element_types,
                    "page_numbers": chunk_page_numbers,
                    "original_text": chunk_text,
                    "is_appendix": is_appendix,  # Phase 22
                    "appendix_id": appendix_id   # Phase 22
                }

                doc = Document(page_content=normalized_text, metadata=metadata)
                documents.append(doc)
                chunk_id += 1

    # Update total_chunks for all documents
    total_chunks = len(documents)
    for doc in documents:
        doc.metadata['total_chunks'] = total_chunks

    logger.info(f"Created {total_chunks} chunks from {len(sections)} sections")
    return documents





def load_docx_document(file_path: Path) -> str:
    """
    Load a DOCX document.

    Args:
        file_path: Path to the DOCX file

    Returns:
        Extracted text content

    Raises:
        ImportError: If DOCX loader is not available
        Exception: If DOCX cannot be read
    """
    try:
        doc = DocxDocument(str(file_path))
        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
        logger.info(f"Extracted {len(paragraphs)} paragraphs from {file_path.name}")
        return "\n\n".join(paragraphs)

    except Exception as e:
        logger.error(f"Failed to load DOCX {file_path.name}: {str(e)}")
        raise


def load_document(file_path: Path) -> str:
    """
    Load a document based on its file extension.

    Args:
        file_path: Path to the document

    Returns:
        Extracted text content

    Raises:
        ValueError: If file format is not supported
        Exception: If document cannot be loaded
    """
    extension = file_path.suffix.lower()

    if extension == '.txt':
        return load_txt_document(file_path)
    elif extension == '.pdf':
        return load_pdf_document(file_path)
    elif extension == '.docx':
        return load_docx_document(file_path)
    else:
        raise ValueError(f"Unsupported file format: {extension}")


# ==============================================================================
# DOCUMENT CHUNKING AND METADATA
# ==============================================================================

def chunk_document(
    text: str,
    document_id: str,
    document_name: str,
    file_type: str,
    ingestion_timestamp: str,
    chunk_size: int = 500,
    chunk_overlap: int = 50,
    file_path: Optional[Path] = None  # Phase 18: Add file_path for layout-aware parsing
) -> List[Document]:
    """
    Chunk a document and attach metadata to each chunk.
    
    Phase 18: PDFs use layout-aware parsing when available, falling back to
    linear chunking if parsing fails or unstructured library is not installed.

    Args:
        text: Document text content
        document_id: Unique document identifier
        document_name: Original filename
        file_type: File extension (.txt, .pdf, .docx)
        ingestion_timestamp: ISO format timestamp
        chunk_size: Maximum chunk size in characters
        chunk_overlap: Overlap between chunks
        file_path: Path to original file (Phase 18, for layout-aware PDF parsing)

    Returns:
        List of Document objects with metadata
    """
    # Phase 18: Try layout-aware parsing for PDFs
    if file_type == '.pdf' and file_path:
        try:
            logger.info(f"Using layout-aware parsing for {document_name}")
            elements = load_pdf_document_layout_aware(file_path)

            # If no elements found, fall back to linear chunking
            if not elements:
                raise ValueError(f"Layout-aware parsing returned 0 elements for {document_name}")

            # Phase 18.1: Text normalization pipeline
            from modules.text_normalizer_pipeline import normalize_elements, get_normalization_stats
            original_element_count = len(elements)
            elements = normalize_elements(elements)
            
            # Log normalization stats
            stats = get_normalization_stats(
                [{'text': ''}] * original_element_count,  # Placeholder for original count
                elements
            )
            logger.info(
                f"Text normalization: {stats['original_element_count']} → {stats['normalized_element_count']} elements "
                f"({stats['elements_removed']} removed)"
            )
            
            sections = group_elements_by_section(elements)
            documents = create_chunks_from_sections(
                sections=sections,
                document_id=document_id,
                document_name=document_name,
                file_type=file_type,
                ingestion_timestamp=ingestion_timestamp,
                target_size=chunk_size,
                max_size=min(chunk_size * 2, 800)  # Max 800 chars or 2x chunk_size
            )
            logger.info(f"Layout-aware parsing completed for {document_name}")
            return documents
        except Exception as e:
            logger.warning(f"Layout-aware parsing failed for {document_name}: {str(e)}")
            logger.info(f"Falling back to linear chunking for {document_name}")
            # Fall through to linear chunking below
    
    # Linear chunking (original Phase 17C logic)
    # Used for: .txt, .docx, PDFs when layout parsing unavailable/fails
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        length_function=len,
        separators=["\n\n", "\n", ". ", " ", ""]
    )

    # Split text into chunks
    chunks = text_splitter.split_text(text)

    # Create Document objects with rich metadata
    # Text normalization: Store original text in metadata, use normalized for embedding
    documents = []
    for i, chunk in enumerate(chunks):
        section_name = extract_section_name(chunk)

        # Normalize text for consistent embedding (case-insensitive matching)
        normalized_chunk = normalize_text(chunk)

        metadata = {
            "document_id": document_id,
            "document_name": document_name,
            "file_type": file_type,
            "ingestion_timestamp": ingestion_timestamp,
            "chunk_id": i,
            "total_chunks": len(chunks),
            "section": section_name,
            "original_text": chunk  # Preserve original text for display/citation
        }

        # Use normalized text for page_content (embedding), original stored in metadata
        doc = Document(page_content=normalized_chunk, metadata=metadata)
        documents.append(doc)

    logger.info(f"Created {len(chunks)} chunks from {document_name}")
    return documents


# ==============================================================================
# DOCUMENT INGESTION
# ==============================================================================

class DocumentManager:
    """
    Main document management class that handles ingestion, indexing, and retrieval.
    """

    def __init__(
        self,
        registry_path: Path,
        vector_store_path: Path,
        openai_api_key: Optional[str] = None
    ):
        """
        Initialize the document manager.

        Args:
            registry_path: Path to the document registry JSON file
            vector_store_path: Path to the vector store directory
            openai_api_key: OpenAI API key (uses env var if not provided)
        """
        self.registry = DocumentRegistry(registry_path)
        self.vector_store_path = vector_store_path
        self.api_key = openai_api_key or OPENAI_API_KEY

        if not self.api_key:
            raise ValueError("OpenAI API key not found. Set OPENAI_API_KEY environment variable.")

        self.embeddings = OpenAIEmbeddings(openai_api_key=self.api_key)
        self.vector_store = None

    def load_vector_store(self) -> Optional[FAISS]:
        """
        Load existing vector store from disk.

        Returns:
            FAISS vector store or None if not found
        """
        if self.vector_store_path.exists():
            logger.info(f"Loading vector store from {self.vector_store_path}")
            self.vector_store = FAISS.load_local(
                str(self.vector_store_path),
                self.embeddings,
                allow_dangerous_deserialization=True
            )
            return self.vector_store
        logger.warning("No existing vector store found.")
        return None

    def save_vector_store(self):
        """Save the vector store to disk."""
        if self.vector_store:
            self.vector_store_path.mkdir(parents=True, exist_ok=True)
            self.vector_store.save_local(str(self.vector_store_path))
            logger.info(f"Vector store saved to {self.vector_store_path}")

    def ingest_document(
        self,
        file_path: Path,
        chunk_size: int = 500,
        chunk_overlap: int = 50,
        skip_duplicates: bool = True
    ) -> Tuple[bool, str]:
        """
        Ingest a single document into the vector store.

        Args:
            file_path: Path to the document file
            chunk_size: Chunk size for text splitting
            chunk_overlap: Overlap between chunks
            skip_duplicates: If True, skip documents that already exist

        Returns:
            Tuple of (success: bool, message: str)
        """
        try:
            # Validate file exists
            if not file_path.exists():
                return False, f"File not found: {file_path}"

            # Validate file format
            if file_path.suffix.lower() not in SUPPORTED_FORMATS:
                return False, f"Unsupported format: {file_path.suffix}"

            # Calculate file hash for duplicate detection
            file_hash = calculate_file_hash(file_path)

            # Check for duplicates
            if skip_duplicates and self.registry.document_exists(file_hash):
                existing = self.registry.get_by_hash(file_hash)
                if existing:
                    doc_id, metadata = existing
                    return False, f"Duplicate: {metadata['document_name']} already ingested on {metadata['ingestion_timestamp']}"

            # Generate document ID and metadata
            document_id = hashlib.md5(f"{file_path.name}{datetime.now().isoformat()}".encode()).hexdigest()
            document_name = file_path.name
            file_type = file_path.suffix.lower()
            ingestion_timestamp = datetime.now().isoformat()

            logger.info(f"Ingesting document: {document_name}")

            # Load document content
            text_content = load_document(file_path)

            if not text_content.strip():
                return False, f"Document is empty: {document_name}"

            # Chunk document with metadata (Phase 18: pass file_path for layout-aware parsing)
            documents = chunk_document(
                text=text_content,
                document_id=document_id,
                document_name=document_name,
                file_type=file_type,
                ingestion_timestamp=ingestion_timestamp,
                chunk_size=chunk_size,
                chunk_overlap=chunk_overlap,
                file_path=file_path  # Phase 18: Enable layout-aware PDF parsing
            )

            # Phase 24: Generic consolidation engine (replaces hardcoded Phase 18/21.1)
            from modules.consolidation_engine import ConsolidationEngine
            engine = ConsolidationEngine()
            documents = engine.consolidate_all(documents)

            # Phase 23: Validate metadata and log statistics
            logger.info("[PHASE23] Validating chunk metadata...")
            is_pdf_layout = file_type == '.pdf'
            validate_all_chunks(documents, is_pdf_layout)
            log_metadata_statistics(documents)

            # Add to vector store (incremental)
            if self.vector_store is None:
                # Create new vector store
                logger.info("Creating new vector store")
                self.vector_store = FAISS.from_documents(documents, self.embeddings)
            else:
                # Add to existing vector store incrementally
                logger.info("Adding documents to existing vector store")
                self.vector_store.add_documents(documents)

            # Save vector store
            self.save_vector_store()

            # Phase 25: Build and save metadata index
            from modules.metadata_index import MetadataIndex
            metadata_index = MetadataIndex()
            metadata_index.build_from_vector_store(self.vector_store)
            metadata_index.save()

            # Add to registry
            registry_metadata = {
                "document_name": document_name,
                "file_type": file_type,
                "file_hash": file_hash,
                "file_path": str(file_path.absolute()),
                "ingestion_timestamp": ingestion_timestamp,
                "num_chunks": len(documents),
                "chunk_size": chunk_size,
                "chunk_overlap": chunk_overlap
            }
            self.registry.add_document(document_id, registry_metadata)

            return True, f"Successfully ingested {document_name} ({len(documents)} chunks)"

        except Exception as e:
            logger.error(f"Error ingesting {file_path.name}: {str(e)}")
            return False, f"Error: {str(e)}"

    def ingest_directory(
        self,
        directory_path: Path,
        chunk_size: int = 500,
        chunk_overlap: int = 50,
        skip_duplicates: bool = True
    ) -> Dict[str, List[str]]:
        """
        Ingest all supported documents from a directory.

        Args:
            directory_path: Path to directory containing documents
            chunk_size: Chunk size for text splitting
            chunk_overlap: Overlap between chunks
            skip_duplicates: If True, skip documents that already exist

        Returns:
            Dictionary with 'success' and 'failed' lists of filenames
        """
        results = {"success": [], "failed": [], "skipped": []}

        if not directory_path.exists():
            logger.error(f"Directory not found: {directory_path}")
            return results

        # Find all supported files
        files = []
        for ext in SUPPORTED_FORMATS:
            files.extend(directory_path.glob(f"*{ext}"))

        if not files:
            logger.warning(f"No supported documents found in {directory_path}")
            return results

        logger.info(f"Found {len(files)} document(s) to ingest")

        # Ingest each file
        for file_path in files:
            success, message = self.ingest_document(
                file_path,
                chunk_size=chunk_size,
                chunk_overlap=chunk_overlap,
                skip_duplicates=skip_duplicates
            )

            if success:
                results["success"].append(file_path.name)
            elif "Duplicate" in message:
                results["skipped"].append(file_path.name)
            else:
                results["failed"].append(file_path.name)

            logger.info(f"{file_path.name}: {message}")

        return results

    def delete_document(self, document_id: str) -> Tuple[bool, str]:
        """
        Delete a document from the vector store and registry using native FAISS deletion.

        Uses LangChain FAISS.delete(ids) to remove chunks by their docstore UUIDs,
        avoiding the need to re-read source files or re-embed.

        Args:
            document_id: Document ID to delete

        Returns:
            Tuple of (success: bool, message: str)
        """
        doc_metadata = self.registry.get_document(document_id)
        if not doc_metadata:
            return False, f"Document not found: {document_id}"

        document_name = doc_metadata.get('document_name', 'Unknown')

        try:
            # Step 1: Collect docstore UUIDs for this document
            uuids_to_delete = []
            if self.vector_store is not None:
                for faiss_idx, docstore_id in self.vector_store.index_to_docstore_id.items():
                    doc = self.vector_store.docstore.search(docstore_id)
                    if doc and doc.metadata.get('document_id') == document_id:
                        uuids_to_delete.append(docstore_id)

            logger.info(f"Found {len(uuids_to_delete)} chunks to delete for '{document_name}' ({document_id})")

            # Step 2: Delete from FAISS vector store
            if uuids_to_delete:
                total_chunks = len(self.vector_store.index_to_docstore_id)
                if len(uuids_to_delete) >= total_chunks:
                    # Deleting the last document — clear vector store entirely
                    logger.info("Deleting last document — clearing vector store")
                    import shutil
                    if self.vector_store_path.exists():
                        shutil.rmtree(self.vector_store_path)
                    self.vector_store = None
                else:
                    # Native FAISS deletion — no re-embedding needed
                    self.vector_store.delete(uuids_to_delete)
                    self.save_vector_store()
            elif self.vector_store is not None:
                logger.warning(f"No chunks found in FAISS for '{document_name}' ({document_id}) — stale registry entry")

            # Step 3: Remove from registry AFTER successful FAISS deletion
            self.registry.remove_document(document_id)

            # Step 4: Rebuild metadata index from updated vector store
            try:
                from modules.metadata_index import MetadataIndex
                metadata_index = MetadataIndex()
                metadata_index.build_from_vector_store(self.vector_store)
                metadata_index.save()
            except Exception as e:
                logger.warning(f"Metadata index rebuild failed (non-fatal): {e}")

            return True, f"Deleted '{document_name}': removed {len(uuids_to_delete)} chunks from vector store"

        except Exception as e:
            logger.error(f"Error deleting document '{document_name}' ({document_id}): {e}")
            return False, f"Deletion failed: {str(e)}"

    def rebuild_vector_store(self) -> Tuple[bool, str]:
        """
        Rebuild the vector store from all documents in the registry.

        This method re-ingests all documents from their original file paths.

        Returns:
            Tuple of (success: bool, message: str)
        """
        try:
            documents_list = self.registry.list_documents()

            if not documents_list:
                # No documents, clear vector store
                if self.vector_store_path.exists():
                    import shutil
                    shutil.rmtree(self.vector_store_path)
                self.vector_store = None
                return True, "Vector store cleared (no documents)"

            all_documents = []

            # Re-load and chunk each document
            for doc_info in documents_list:
                file_path = Path(doc_info['file_path'])

                if not file_path.exists():
                    logger.warning(f"File not found, skipping: {file_path}")
                    continue

                # Load document
                text_content = load_document(file_path)

                # Chunk with original metadata (Phase 18: pass file_path for layout-aware parsing)
                documents = chunk_document(
                    text=text_content,
                    document_id=doc_info['document_id'],
                    document_name=doc_info['document_name'],
                    file_type=doc_info['file_type'],
                    ingestion_timestamp=doc_info['ingestion_timestamp'],
                    chunk_size=doc_info.get('chunk_size', 500),
                    chunk_overlap=doc_info.get('chunk_overlap', 50),
                    file_path=file_path  # Phase 18: Enable layout-aware PDF parsing
                )

                all_documents.extend(documents)

            # Rebuild vector store
            logger.info(f"Rebuilding vector store with {len(all_documents)} chunks from {len(documents_list)} documents")

            # Handle edge case: no chunks to rebuild
            if len(all_documents) == 0:
                logger.warning("No chunks to rebuild - clearing vector store")
                # Clear vector store files
                vector_store_path = Path(__file__).parent / "vector_store"
                if vector_store_path.exists():
                    import shutil
                    shutil.rmtree(vector_store_path)
                self.vector_store = None
                return True, "All documents removed - vector store cleared"

            self.vector_store = FAISS.from_documents(all_documents, self.embeddings)
            self.save_vector_store()

            # Phase 25: Rebuild metadata index
            from modules.metadata_index import MetadataIndex
            metadata_index = MetadataIndex()
            metadata_index.build_from_vector_store(self.vector_store)
            metadata_index.save()

            return True, f"Rebuilt vector store with {len(documents_list)} documents"

        except Exception as e:
            logger.error(f"Error rebuilding vector store: {str(e)}")
            return False, str(e)

    def list_documents(self) -> List[Dict]:
        """
        List all ingested documents.

        Returns:
            List of document metadata dictionaries
        """
        return self.registry.list_documents()

    def get_document_info(self, document_id: str) -> Optional[Dict]:
        """
        Get information about a specific document.

        Args:
            document_id: Document identifier

        Returns:
            Document metadata or None
        """
        return self.registry.get_document(document_id)
